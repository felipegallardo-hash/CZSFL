from __future__ import annotations

import io
import math
from datetime import date
from decimal import Decimal, ROUND_HALF_UP
from pathlib import Path

import pandas as pd
import streamlit as st
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

APP_DIR = Path(__file__).resolve().parent
DATA_FILE = APP_DIR / "Cotizador_Empaques_2026_VERSION_FINAL.xlsx"
WORD_TEMPLATE = APP_DIR / "Plantilla.docx"

st.set_page_config(
    page_title="Cotizador Sellado de Fluidos",
    page_icon="🧾",
    layout="wide",
)

# ---------- Styling ----------
st.markdown(
    """
    <style>
    .block-container {padding-top: 1.5rem; padding-bottom: 2rem; max-width: 1250px;}
    .main-title {
        background:#17365D; color:white; padding:18px 22px; border-radius:12px;
        font-size:30px; font-weight:700; margin-bottom:8px;
    }
    .sub-title {color:#666; margin-bottom:20px;}
    .result-box {
        background:#E2F0D9; border:1px solid #B7C9D6; border-radius:10px;
        padding:14px 16px; margin-bottom:10px;
    }
    .cost-box {
        background:#D9EAF7; border:1px solid #B7C9D6; border-radius:10px;
        padding:14px 16px; margin-bottom:10px;
    }
    .small-muted {font-size:0.9rem; color:#6b7280;}
    </style>
    """,
    unsafe_allow_html=True,
)

def clp(value: float | int | None) -> str:
    if value is None:
        return "-"
    return "$ " + f"{round(float(value)):,.0f}".replace(",", ".")

def excel_round_hundred(value: float) -> int:
    """Equivalent to Excel ROUND(value,-2) for positive quote values."""
    d = Decimal(str(value)) / Decimal("100")
    return int(d.quantize(Decimal("1"), rounding=ROUND_HALF_UP) * Decimal("100"))

def excel_roundup_integer(value: float) -> int:
    """Equivalent to Excel ROUNDUP(value,0) for positive quote values."""
    return int(math.ceil(value))

def sale_price_from_factor(cost: float, factor: float) -> int:
    """Cotizador original: ROUND(costo * factor, -2)."""
    return excel_round_hundred(cost * factor)

def sale_price_from_margin(cost: float, margin_pct: float) -> int:
    """Cotizador original: ROUNDUP(costo / (1 - margen), 0)."""
    margin = margin_pct / 100.0
    if margin >= 1:
        raise ValueError("El margen debe ser menor a 100%.")
    return excel_roundup_integer(cost / (1 - margin))

@st.cache_data(show_spinner=False)
def load_master_data(path: str):
    wb_values = load_workbook(path, data_only=True, read_only=True)

    # ---- Flat gaskets / cuts ----
    listado = wb_values["LISTADO"]

    cuts = {}
    for row in listado.iter_rows(min_row=2, min_col=1, max_col=12, values_only=True):
        key = row[0]       # CONCATENADO, e.g. FF150 3-1/2"
        if not key:
            continue
        key = str(key)
        # Excel VLOOKUP(...,FALSE) returns the FIRST exact match.
        # Do not overwrite a value if the same key appears later in LISTADO.
        if key not in cuts:
            cuts[key] = {
                "od_mm": float(row[3]) if row[3] is not None else None,
                "id_mm": float(row[4]) if row[4] is not None else None,
                "class": row[7],
                "cut_cost": float(row[11]) if row[11] is not None else 0.0,
            }

    materials = {}
    for row in listado.iter_rows(min_row=2, min_col=33, max_col=37, values_only=True):
        name, inventory, unit_cost, length, width = row
        if not name or length is None or width is None:
            continue
        name = str(name)
        # Excel VLOOKUP(...,FALSE) returns the FIRST exact match.
        # LISTADO contains repeated material names; later duplicates must NOT overwrite
        # the first record used by the original cotizador.
        if name not in materials:
            materials[name] = {
                "inventory": float(inventory) if inventory is not None else 0.0,
                "sheet_cost": float(unit_cost) if unit_cost is not None else 0.0,
                "length": float(length),
                "width": float(width),
            }

    # ---- Braided packing ----
    # LISTADO AC=style, AD=section, AF=real product code.
    braid_map = []
    for r in range(14, 30):
        style = listado.cell(r, 29).value
        section = listado.cell(r, 30).value
        code = listado.cell(r, 32).value
        if style is not None and section is not None and code:
            braid_map.append((style, str(section), str(code)))

    mayor = wb_values["MAYOR AUX G"]
    cost_inventory = {}
    for row in mayor.iter_rows(min_row=3, min_col=1, max_col=11, values_only=True):
        code = row[0]
        if not code:
            continue
        cost_inventory[str(code).strip("'")] = {
            "description": row[1],
            "stock": float(row[5]) if row[5] is not None else 0.0,
            "cost": float(row[10]) if row[10] is not None else 0.0,
        }

    braids = {}
    for style, section, code in braid_map:
        source = cost_inventory.get(code, {})
        braids[(str(style), section)] = {
            "code": code,
            "cost": float(source.get("cost", 0.0)),
            "stock": float(source.get("stock", 0.0)),
            "description": source.get("description", ""),
        }

    return cuts, materials, braids


def _replace_text_in_paragraph(paragraph, replacements: dict[str, str]):
    """Replace visible text while preserving the paragraph's first-run formatting."""
    original = paragraph.text
    updated = original
    for old, new in replacements.items():
        if old in updated:
            updated = updated.replace(old, new)
    if updated != original:
        if paragraph.runs:
            paragraph.runs[0].text = updated
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.text = updated

def _replace_text_everywhere(doc, replacements: dict[str, str]):
    for p in doc.paragraphs:
        _replace_text_in_paragraph(p, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_text_in_paragraph(p, replacements)

def _set_cell_text(cell, text, bold=False, size=8, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    r = p.add_run(str(text))
    r.bold = bold
    r.font.size = Pt(size)

def quote_docx(
    items: list[dict],
    client: str,
    quote_no: str,
    quote_date: date,
    attention: str,
    phone: str,
    email: str,
    address: str,
    city: str,
    reference: str,
    currency: str,
    payment_terms: str,
    delivery_place: str,
    delivery_days: str,
    validity_days: str,
    scope: str,
) -> bytes:
    """Fill the corporate Invenio Word template and return a DOCX."""
    doc = Document(str(WORD_TEMPLATE))

    # Header / commercial data.
    replacements = {
        "SFL0826-CZ0014": quote_no or "S/N",
        "18-08-26": quote_date.strftime("%d-%m-%y"),
        "Referencia G": f"Referencia {reference}".rstrip(),
        "Pesos chilenos": currency or "Pesos chilenos",
        "Facturación 30 Días": payment_terms or "Facturación 30 Días",
        "Bodega de Transito Santiago": delivery_place or "Bodega de Transito Santiago",
        "15 Días": f"{delivery_days} Días" if delivery_days else "15 Días",
        "Validez Oferta (días) 10": f"Validez Oferta (días) {validity_days or '10'}",
        "Señor (es):": f"Señor (es): {client}".rstrip(),
        "At. Sr.(es/ta)": f"At. Sr.(es/ta) {attention}".rstrip(),
        "Fono": f"Fono {phone}".rstrip(),
        "e-mail": f"e-mail {email}".rstrip(),
        "Dirección": f"Dirección {address}".rstrip(),
        "Ciudad": f"Ciudad {city}".rstrip(),
    }
    _replace_text_everywhere(doc, replacements)

    # Identify the item table by its header row.
    item_table = None
    for table in doc.tables:
        if not table.rows:
            continue
        header = " | ".join(cell.text for cell in table.rows[0].cells)
        if "Ítem" in header and "Descripción" in header and "Valor unitario" in header:
            item_table = table
            break

    if item_table is not None:
        # The template has: header row, one item row, total-label row, red-note row,
        # and an Alcances area below. We fill the existing first item row and insert
        # additional rows immediately after it, preserving the table style.
        base_row = item_table.rows[1]

        def fill_item_row(row, idx, item):
            qty = item["Cantidad"]
            if isinstance(qty, float) and qty.is_integer():
                qty = int(qty)
            values = [
                f"{idx}.",
                qty,
                item["Glosa"],
                item.get("Unidad", ""),
                clp(item["Precio unitario"]).replace("$ ", "$"),
                clp(item["Subtotal"]).replace("$ ", "$"),
            ]
            for c, value in enumerate(values):
                align = WD_ALIGN_PARAGRAPH.CENTER if c in (0,1,3,4,5) else WD_ALIGN_PARAGRAPH.LEFT
                _set_cell_text(row.cells[c], value, bold=(c == 0), size=8, align=align)

        if items:
            fill_item_row(base_row, 1, items[0])

            # Insert extra product rows before the summary row.
            for idx, item in enumerate(items[1:], 2):
                new_row = item_table.add_row()
                # Move newly-added row before the current summary row (row index 2 initially).
                tr = new_row._tr
                summary_tr = item_table.rows[2]._tr
                summary_tr.addprevious(tr)
                fill_item_row(new_row, idx, item)
        else:
            fill_item_row(base_row, 1, {
                "Cantidad": "",
                "Glosa": "",
                "Unidad": "",
                "Precio unitario": 0,
                "Subtotal": 0,
            })

        total = sum(float(x["Subtotal"]) for x in items)
        # Find the summary row containing PRECIOS UNITARIOS NETOS and place total in last cell.
        for row in item_table.rows:
            joined = " ".join(c.text for c in row.cells)
            if "PRECIOS UNITARIOS NETOS" in joined:
                _set_cell_text(
                    row.cells[-1],
                    clp(total).replace("$ ", "$"),
                    bold=True,
                    size=8,
                    align=WD_ALIGN_PARAGRAPH.RIGHT,
                )
                break

        # Alcances: write scope into the cell/paragraph containing the label.
        if scope:
            for row in item_table.rows:
                joined = " ".join(c.text for c in row.cells)
                if "Alcances" in joined:
                    cell = row.cells[0]
                    if len(cell.paragraphs) == 1:
                        cell.add_paragraph(scope)
                    else:
                        cell.paragraphs[-1].text = scope
                    break

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

def quote_xlsx(items: list[dict], client: str, quote_no: str, quote_date: date) -> bytes:
    wb = Workbook()
    ws = wb.active
    ws.title = "COTIZACION"

    navy = "17365D"
    green = "E2F0D9"
    white = "FFFFFF"
    thin = Side(style="thin", color="B7C9D6")

    ws.merge_cells("A1:F1")
    ws["A1"] = "COTIZACIÓN DE EMPAQUETADURAS"
    ws["A1"].font = Font(size=16, bold=True, color=white)
    ws["A1"].fill = PatternFill("solid", fgColor=navy)
    ws["A1"].alignment = Alignment(horizontal="center")

    ws["A3"], ws["B3"] = "Cliente", client
    ws["D3"], ws["E3"] = "N° Cotización", quote_no
    ws["A4"], ws["B4"] = "Fecha", quote_date.strftime("%d-%m-%Y")

    headers = ["N°", "Tipo", "Glosa", "Cantidad", "Precio unit.", "Subtotal"]
    for c, h in enumerate(headers, 1):
        cell = ws.cell(6, c, h)
        cell.font = Font(bold=True, color=white)
        cell.fill = PatternFill("solid", fgColor=navy)
        cell.alignment = Alignment(horizontal="center")
        cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)

    for i, item in enumerate(items, 1):
        r = 6 + i
        values = [
            i,
            item["Tipo"],
            item["Glosa"],
            item["Cantidad"],
            item["Precio unitario"],
            item["Subtotal"],
        ]
        for c, v in enumerate(values, 1):
            cell = ws.cell(r, c, v)
            cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)
            cell.alignment = Alignment(vertical="top", wrap_text=True)
        ws.cell(r, 5).number_format = '$ #,##0'
        ws.cell(r, 6).number_format = '$ #,##0'

    total_row = 8 + len(items)
    ws.cell(total_row, 5, "TOTAL")
    ws.cell(total_row, 5).font = Font(bold=True, color=white)
    ws.cell(total_row, 5).fill = PatternFill("solid", fgColor=navy)
    ws.cell(total_row, 6, sum(x["Subtotal"] for x in items))
    ws.cell(total_row, 6).font = Font(bold=True)
    ws.cell(total_row, 6).fill = PatternFill("solid", fgColor=green)
    ws.cell(total_row, 6).number_format = '$ #,##0'

    widths = {"A": 7, "B": 18, "C": 55, "D": 14, "E": 18, "F": 18}
    for col, width in widths.items():
        ws.column_dimensions[col].width = width

    buffer = io.BytesIO()
    wb.save(buffer)
    return buffer.getvalue()

cuts, materials, braids = load_master_data(str(DATA_FILE))

if "quote_items" not in st.session_state:
    st.session_state.quote_items = []

st.markdown('<div class="main-title">COTIZADOR · SELLADO DE FLUIDOS</div>', unsafe_allow_html=True)
st.markdown(
    '<div class="sub-title">Cotizador de empaquetaduras planas y trenzadas basado en la matriz de costos validada.</div>',
    unsafe_allow_html=True,
)
st.info(
    "Importante: usa **Factor multiplicador** cuando trabajes con valores como 2 o 6,5. "
    "Usa **Margen %** cuando quieras ingresar, por ejemplo, 50% de margen. "
    "La app aplica la misma fórmula del cotizador original para cada caso."
)

st.warning(
    "Los ajustes manuales de costo se aplican a la cotización actual. "
    "El costo base del Excel no se modifica automáticamente."
)

# Quote metadata
m1, m2, m3 = st.columns([2.2, 1.2, 1])
with m1:
    client = st.text_input("Cliente", placeholder="Ej.: Compañía Minera...")
with m2:
    quote_no = st.text_input("N° Cotización", placeholder="SFL0826-CZ...")
with m3:
    quote_date = st.date_input("Fecha", value=date.today())

with st.expander("Datos para la plantilla Word", expanded=False):
    d1, d2, d3 = st.columns(3)
    with d1:
        attention = st.text_input("Atención")
        phone = st.text_input("Fono")
        email = st.text_input("e-mail")
        address = st.text_input("Dirección")
        city = st.text_input("Ciudad")
    with d2:
        reference = st.text_input("Referencia")
        currency = st.text_input("Precios", value="Pesos chilenos")
        payment_terms = st.text_input("Forma de pago", value="Facturación 30 Días")
    with d3:
        delivery_place = st.text_input("Materiales puesto en", value="Bodega de Transito Santiago")
        delivery_days = st.text_input("Plazo de entrega (días)", value="15")
        validity_days = st.text_input("Validez oferta (días)", value="10")
    scope = st.text_area("Alcances", placeholder="Opcional")

tab_flat, tab_braid, tab_quote = st.tabs(
    ["🟦 Empaquetadura plana", "🟩 Empaquetadura trenzada", "🧾 Cotización"]
)

# ========================= FLAT =========================
with tab_flat:
    st.subheader("Empaquetadura plana / cortada")
    c1, c2 = st.columns(2)

    with c1:
        material = st.selectbox("Material", sorted(materials.keys()))
        cut = st.selectbox("Tipo de corte", list(cuts.keys()))
        qty = st.number_input("Cantidad a cotizar", min_value=1, value=1, step=1)

    with c2:
        pricing_mode_flat = st.radio(
            "Método de precio",
            ["Factor multiplicador", "Margen %"],
            horizontal=True,
            key="pricing_mode_flat",
            help="El cotizador original usa fórmulas distintas para factor y margen.",
        )
        if pricing_mode_flat == "Factor multiplicador":
            factor = st.number_input(
                "Factor",
                min_value=0.01,
                value=2.0,
                step=0.1,
                format="%.2f",
                key="flat_factor",
                help="Ejemplo: 2,0 duplica el costo.",
            )
            margin_pct_flat = None
        else:
            margin_pct_flat = st.number_input(
                "Margen (%)",
                min_value=0.0,
                max_value=99.9,
                value=50.0,
                step=1.0,
                format="%.1f",
                key="flat_margin",
                help="Ejemplo: para 50% ingresa 50.",
            )
            factor = None

    mat = materials[material]
    cut_data = cuts[cut]

    od = cut_data["od_mm"]
    pieces = 0
    if od and od > 0:
        pieces = math.floor(mat["length"] / od) * math.floor(mat["width"] / od)

    cut_cost = cut_data["cut_cost"]

    st.markdown("#### Ajuste de costo de materia prima")
    st.caption(
        "El valor mostrado corresponde al costo base del Excel. "
        "Puedes modificarlo para reflejar alzas o disminuciones sin cambiar la base original."
    )
    raw_sheet_cost = st.number_input(
        "Costo actual de la plancha",
        min_value=0.0,
        value=float(mat["sheet_cost"]),
        step=1000.0,
        format="%.2f",
        key=f"flat_raw_cost_{material}",
    )

    material_unit_cost = (raw_sheet_cost / pieces) if pieces else 0.0
    flat_cost = cut_cost + material_unit_cost
    if pieces:
        if pricing_mode_flat == "Factor multiplicador":
            unit_sale = sale_price_from_factor(flat_cost, factor)
        else:
            unit_sale = sale_price_from_margin(flat_cost, margin_pct_flat)
    else:
        unit_sale = 0
    total = unit_sale * qty
    glosa = f"EMPAQUETADURA {cut} {material}"

    r1, r2, r3, r4 = st.columns(4)
    r1.metric("Cantidad por plancha", f"{pieces}")
    r2.metric("Costo corte unit.", clp(cut_cost))
    r3.metric("Costo material / un.", clp(material_unit_cost))
    delta_raw = raw_sheet_cost - float(mat["sheet_cost"])
    r4.metric(
        "Costo plancha ajustado",
        clp(raw_sheet_cost),
        delta=clp(delta_raw) if abs(delta_raw) >= 0.5 else None,
    )

    p1, p2 = st.columns(2)
    with p1:
        st.markdown(
            f'<div class="result-box"><b>Precio venta unitario</b><br><span style="font-size:1.65rem">{clp(unit_sale)}</span></div>',
            unsafe_allow_html=True,
        )
    with p2:
        st.markdown(
            f'<div class="result-box"><b>Total cotización</b><br><span style="font-size:1.65rem">{clp(total)}</span></div>',
            unsafe_allow_html=True,
        )

    st.caption(
        f"{glosa} · Costo plancha base: {clp(mat['sheet_cost'])} · "
        f"Costo plancha usado: {clp(raw_sheet_cost)} · OD de corte: {cut_data['od_mm']:.0f} mm"
    )

    if st.button("➕ Agregar plana a cotización", type="primary", use_container_width=True):
        st.session_state.quote_items.append(
            {
                "Tipo": "PLANA",
                "Glosa": glosa,
                "Cantidad": qty,
                "Unidad": "UN",
                "Precio unitario": unit_sale,
                "Subtotal": total,
                "Costo materia prima": raw_sheet_cost,
                "Costo base materia prima": mat["sheet_cost"],
            }
        )
        st.success("Producto agregado a la cotización.")

# ========================= BRAIDED =========================
with tab_braid:
    st.subheader("Empaquetadura trenzada")

    styles = list(dict.fromkeys(key[0] for key in braids.keys()))
    b1, b2 = st.columns(2)

    with b1:
        style = st.selectbox("Estilo", styles)
        available_sections = [key[1] for key in braids.keys() if key[0] == style]
        section = st.selectbox("Sección", available_sections)
        kg = st.number_input("Kg a cotizar", min_value=0.01, value=2.0, step=0.5, format="%.2f")

    with b2:
        pricing_mode_braid = st.radio(
            "Método de precio",
            ["Factor multiplicador", "Margen %"],
            horizontal=True,
            key="pricing_mode_braid",
            help="El cotizador original usa fórmulas distintas para factor y margen.",
        )
        if pricing_mode_braid == "Factor multiplicador":
            braid_factor = st.number_input(
                "Factor",
                min_value=0.01,
                value=6.5,
                step=0.1,
                format="%.2f",
                key="braid_factor",
                help="Ejemplo: 6,5 multiplica el costo por kg por 6,5.",
            )
            braid_margin_pct = None
        else:
            braid_margin_pct = st.number_input(
                "Margen (%)",
                min_value=0.0,
                max_value=99.9,
                value=50.0,
                step=1.0,
                format="%.1f",
                key="braid_margin",
                help="Ejemplo: para 50% ingresa 50.",
            )
            braid_factor = None

    braid = braids[(style, section)]

    st.markdown("#### Ajuste de costo de materia prima")
    st.caption(
        "El costo por kg se carga desde el Excel. Puedes modificarlo temporalmente "
        "para cotizar con un costo actualizado."
    )
    base_cost_kg = float(braid["cost"])
    cost_kg = st.number_input(
        "Costo actual por kg",
        min_value=0.0,
        value=base_cost_kg,
        step=100.0,
        format="%.2f",
        key=f"braid_raw_cost_{style}_{section}",
    )

    if pricing_mode_braid == "Factor multiplicador":
        price_kg = sale_price_from_factor(cost_kg, braid_factor)
    else:
        price_kg = sale_price_from_margin(cost_kg, braid_margin_pct)
    braid_total = price_kg * kg
    braid_glosa = f'EMPAQUETADURA TRENZADA ESTILO {style} SECCION {section}'
    display_section = section.replace('"', '')

    a1, a2, a3, a4 = st.columns(4)
    a1.metric("Costo base / kg", clp(base_cost_kg))
    delta_cost = cost_kg - base_cost_kg
    a2.metric(
        "Costo usado / kg",
        clp(cost_kg),
        delta=clp(delta_cost) if abs(delta_cost) >= 0.5 else None,
    )
    a3.metric("Stock registrado (kg)", f"{braid['stock']:,.2f}".replace(",", "."))
    a4.metric("Código", braid["code"])

    p1, p2 = st.columns(2)
    with p1:
        st.markdown(
            f'<div class="result-box"><b>Precio venta / kg</b><br><span style="font-size:1.65rem">{clp(price_kg)}</span></div>',
            unsafe_allow_html=True,
        )
    with p2:
        st.markdown(
            f'<div class="result-box"><b>Total cotización</b><br><span style="font-size:1.65rem">{clp(braid_total)}</span></div>',
            unsafe_allow_html=True,
        )

    st.caption(f"Estilo {style} · Sección {display_section} · {braid['code']}")

    if st.button("➕ Agregar trenzada a cotización", type="primary", use_container_width=True):
        st.session_state.quote_items.append(
            {
                "Tipo": "TRENZADA",
                "Glosa": braid_glosa,
                "Cantidad": kg,
                "Unidad": "KG",
                "Precio unitario": price_kg,
                "Subtotal": braid_total,
                "Costo materia prima": cost_kg,
                "Costo base materia prima": base_cost_kg,
            }
        )
        st.success("Producto agregado a la cotización.")

# ========================= QUOTE =========================
with tab_quote:
    st.subheader("Resumen de cotización")

    if not st.session_state.quote_items:
        st.info("Todavía no has agregado productos.")
    else:
        df = pd.DataFrame(st.session_state.quote_items)
        display_df = df[["Tipo", "Glosa", "Cantidad", "Unidad", "Precio unitario", "Subtotal"]].copy()
        st.dataframe(
            display_df,
            use_container_width=True,
            hide_index=True,
            column_config={
                "Precio unitario": st.column_config.NumberColumn(format="$ %d"),
                "Subtotal": st.column_config.NumberColumn(format="$ %d"),
            },
        )

        quote_total = int(df["Subtotal"].sum())
        st.markdown(
            f'<div class="result-box"><b>TOTAL GENERAL</b><br><span style="font-size:1.9rem">{clp(quote_total)}</span></div>',
            unsafe_allow_html=True,
        )

        c_word, c_excel, c_clear = st.columns(3)

        with c_word:
            docx_bytes = quote_docx(
                st.session_state.quote_items,
                client=client,
                quote_no=quote_no,
                quote_date=quote_date,
                attention=attention,
                phone=phone,
                email=email,
                address=address,
                city=city,
                reference=reference,
                currency=currency,
                payment_terms=payment_terms,
                delivery_place=delivery_place,
                delivery_days=delivery_days,
                validity_days=validity_days,
                scope=scope,
            )
            word_filename = f"Cotizacion_{quote_no or 'Sellado_Fluidos'}.docx".replace("/", "-")
            st.download_button(
                "📄 Descargar cotización Word",
                data=docx_bytes,
                file_name=word_filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                type="primary",
            )

        with c_excel:
            xlsx_bytes = quote_xlsx(
                st.session_state.quote_items,
                client=client,
                quote_no=quote_no,
                quote_date=quote_date,
            )
            filename = f"Cotizacion_{quote_no or 'Sellado_Fluidos'}.xlsx".replace("/", "-")
            st.download_button(
                "⬇️ Descargar Excel",
                data=xlsx_bytes,
                file_name=filename,
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True,
            )

        with c_clear:
            if st.button("🗑️ Vaciar cotización", use_container_width=True):
                st.session_state.quote_items = []
                st.rerun()

st.divider()
st.markdown(
    '<div class="small-muted">Base de costos: Cotizador_Empaques_2026_VERSION_FINAL.xlsx. '
    'Para actualizar los costos de la app, reemplaza ese archivo por una versión actualizada manteniendo el mismo nombre y estructura.</div>',
    unsafe_allow_html=True,
)
